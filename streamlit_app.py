import streamlit as st
from groq import Groq
from sentence_transformers import SentenceTransformer, util
import fitz 
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

#Load API key 
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

if not GROQ_API_KEY:
    st.error("GROQ_API_KEY not found in .env file!")
    st.stop()

client = Groq(api_key=GROQ_API_KEY)
embedder = SentenceTransformer('all-MiniLM-L6-v2')

#Initialize session state for results & tone 
if 'result' not in st.session_state:
    st.session_state.result = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'tone' not in st.session_state:
    st.session_state.tone = "Professional & Confident (default)"

# Page title & layout
st.title("JobAI – AI Resume Optimizer & Cover Letter Generator")

# Clear button
if st.button("Clear", type="secondary"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

#Inputs
uploaded_file = st.file_uploader("Upload your Resume (PDF)", type="pdf")
job_desc = st.text_area("Paste the full Job Description", height=180)

# Tone selector (for regenerate)
tone_options = [
    "Professional & Confident (default)",
    "More Confident & Assertive",
    "Shorter & Concise",
    "Pidgin-friendly & Approachable"
]
selected_tone = st.selectbox("Cover Letter Tone", tone_options, index=tone_options.index(st.session_state.tone))

# Disable main button if inputs missing
analyze_disabled = not (uploaded_file is not None and job_desc.strip())

if st.button("Analyze & Generate / Regenerate", 
             disabled=analyze_disabled, 
             type="primary", 
             use_container_width=True):

    # Only re-process if new upload or first time
    if 'last_file_id' not in st.session_state or st.session_state.last_file_id != uploaded_file.file_id:
        with st.spinner("Reading and extracting your resume..."):
            pdf_bytes = uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            resume_text = ""
            for page in doc:
                resume_text += page.get_text("text") + "\n"
            resume_text = resume_text.strip()
            st.session_state.resume_text = resume_text
            st.session_state.last_file_id = uploaded_file.file_id  # track file

    resume_text = st.session_state.resume_text

    # Quick preview
    st.subheader("Resume Preview (first 600 characters)")
    st.text(resume_text[:600] + "..." if len(resume_text) > 600 else resume_text)

    # ATS similarity
    with st.status("Comparing resume with job description...", expanded=False) as status:
        resume_emb = embedder.encode(resume_text)
        jd_emb = embedder.encode(job_desc)
        similarity = util.cos_sim(resume_emb, jd_emb)[0][0].item()
        st.metric("Initial ATS Match Score", f"{similarity:.0%}")
        status.update(label="Comparison finished!", state="complete", expanded=False)

    # Groq generation with tone
    with st.spinner(f"Generating with tone '{selected_tone}'... usually 5–20 seconds"):
        tone_instruction = ""
        if "Confident & Assertive" in selected_tone:
            tone_instruction = "Use a more assertive, bold, and self-assured tone."
        elif "Shorter & Concise" in selected_tone:
            tone_instruction = "Make the cover letter shorter and more concise (250-350 words)."
        elif "Pidgin-friendly" in selected_tone:
            tone_instruction = "Use simple, approachable language with light Nigerian Pidgin influences where natural."

        prompt = f"""
You are a top Nigerian career coach helping beat ATS in banks, fintech (e.g. Paystack, Flutterwave), and global remote jobs.

Resume:
{resume_text[:5000]}

Job Description:
{job_desc}

{tone_instruction}

Output strictly in Markdown:

## ATS Insights
- Match % (estimate): ...
- Recommended keywords to add (8-12, Nigeria-relevant): ...

## Improved Bullet Points
Rewrite 4-6 key bullets stronger & JD-aligned.

## Tailored Cover Letter
Full letter following the tone above. Greeting, Intro, 2 Body paras, Closing.
"""

        try:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1800,
                temperature=0.6
            )
            result = response.choices[0].message.content
            st.session_state.result = result
            st.session_state.tone = selected_tone

        except Exception as e:
            st.error(f"Error: {str(e)}")
            st.info("Try again or shorten inputs.")

# Display results if available
if st.session_state.result:
    st.markdown("### Your Optimized Results")
    st.markdown(st.session_state.result)

    #  Extract sections for downloads─
    lines = st.session_state.result.split("\n")
    cover_letter = ""
    bullets_section = ""
    in_cover = False
    in_bullets = False

    for line in lines:
        if "## Tailored Cover Letter" in line:
            in_cover = True
            in_bullets = False
            continue
        if "## Improved Bullet Points" in line:
            in_bullets = True
            in_cover = False
            continue
        if in_cover and line.strip():
            cover_letter += line + "\n"
        if in_bullets and line.strip():
            bullets_section += line + "\n"

    # Download Cover Letter as .docx
    def create_docx(text, title="Cover Letter"):
        doc = Document()
        doc.add_heading(title, 0)
        for para in text.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="Download Cover Letter (.docx)",
            data=create_docx(cover_letter, "Tailored Cover Letter"),
            file_name="Tailored_Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_cover"
        )

    with col2:
        st.download_button(
            label="Download Improved Bullets (.txt)",
            data=bullets_section.strip().encode("utf-8"),
            file_name="Improved_Resume_Bullets.txt",
            mime="text/plain",
            key="dl_bullets"
        )

# Footer tip
st.markdown("---")
st.caption("Tip: Change tone and click 'Analyze & Generate' to get new versions without re-uploading.")